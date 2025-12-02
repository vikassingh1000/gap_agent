"""
Parser for extracting text from downloaded documents (PDF, DOCX, Excel)
"""
import os
from pathlib import Path
from typing import Dict, List, Any, Optional


class DocumentParser:
    """Parse documents and extract relevant text"""
    
    def __init__(self):
        """Initialize document parser"""
        self.supported_formats = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv']
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document and extract text
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {'error': 'File not found', 'file_path': str(file_path)}
        
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        elif file_ext == '.csv':
            return self._parse_csv(file_path)
        else:
            return {'error': f'Unsupported file format: {file_ext}', 'file_path': str(file_path)}
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF file"""
        try:
            import PyPDF2
            from pdfplumber import PDF
            
            # Try pdfplumber first (better for text extraction)
            try:
                with open(file_path, 'rb') as f:
                    pdf = PDF(f)
                    text_parts = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    full_text = '\n\n'.join(text_parts)
                    
                    return {
                        'file_path': str(file_path),
                        'file_type': 'pdf',
                        'text': full_text,
                        'page_count': len(pdf.pages),
                        'status': 'success'
                    }
            except ImportError:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text_parts = []
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    full_text = '\n\n'.join(text_parts)
                    
                    return {
                        'file_path': str(file_path),
                        'file_type': 'pdf',
                        'text': full_text,
                        'page_count': len(pdf_reader.pages),
                        'status': 'success'
                    }
        except ImportError:
            return {
                'file_path': str(file_path),
                'file_type': 'pdf',
                'error': 'PDF parsing libraries not installed. Install with: pip install PyPDF2 pdfplumber',
                'status': 'error'
            }
        except Exception as e:
            return {
                'file_path': str(file_path),
                'file_type': 'pdf',
                'error': str(e),
                'status': 'error'
            }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = '\n\n'.join(text_parts)
            
            return {
                'file_path': str(file_path),
                'file_type': 'docx',
                'text': full_text,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'status': 'success'
            }
        except ImportError:
            return {
                'file_path': str(file_path),
                'file_type': 'docx',
                'error': 'python-docx not installed. Install with: pip install python-docx',
                'status': 'error'
            }
        except Exception as e:
            return {
                'file_path': str(file_path),
                'file_type': 'docx',
                'error': str(e),
                'status': 'error'
            }
    
    def _parse_excel(self, file_path: Path) -> Dict[str, Any]:
        """Parse Excel file"""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            all_text = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                # Convert DataFrame to text representation
                sheet_text = df.to_string(index=False)
                sheets_data[sheet_name] = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'text': sheet_text
                }
                all_text.append(f"Sheet: {sheet_name}\n{sheet_text}")
            
            full_text = '\n\n'.join(all_text)
            
            return {
                'file_path': str(file_path),
                'file_type': 'excel',
                'text': full_text,
                'sheets': sheets_data,
                'sheet_count': len(excel_file.sheet_names),
                'status': 'success'
            }
        except ImportError:
            return {
                'file_path': str(file_path),
                'file_type': 'excel',
                'error': 'pandas and openpyxl not installed. Install with: pip install pandas openpyxl',
                'status': 'error'
            }
        except Exception as e:
            return {
                'file_path': str(file_path),
                'file_type': 'excel',
                'error': str(e),
                'status': 'error'
            }
    
    def _parse_csv(self, file_path: Path) -> Dict[str, Any]:
        """Parse CSV file"""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            text = df.to_string(index=False)
            
            return {
                'file_path': str(file_path),
                'file_type': 'csv',
                'text': text,
                'rows': len(df),
                'columns': len(df.columns),
                'status': 'success'
            }
        except ImportError:
            return {
                'file_path': str(file_path),
                'file_type': 'csv',
                'error': 'pandas not installed. Install with: pip install pandas',
                'status': 'error'
            }
        except Exception as e:
            return {
                'file_path': str(file_path),
                'file_type': 'csv',
                'error': str(e),
                'status': 'error'
            }
    
    def parse_directory(self, directory: str) -> List[Dict[str, Any]]:
        """
        Parse all supported documents in a directory
        
        Args:
            directory: Directory path
            
        Returns:
            List of parsed document dictionaries
        """
        directory = Path(directory)
        if not directory.exists():
            return []
        
        results = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                print(f"Parsing: {file_path.name}")
                result = self.parse_document(str(file_path))
                results.append(result)
        
        return results

